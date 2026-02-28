"""Dummy data generation script for the RAG system.

This script programmatically generates realistic dummy documents in multiple formats:
- 5 TXT files on general knowledge topics
- 3 PDF files
- 2 DOCX files
- 2 HTML files

All content is meaningful and realistic, suitable for testing the RAG pipeline.
"""

import os
from pathlib import Path

from docx import Document
from fpdf import FPDF


class ScienceTextFile:
    """Generate science-related text content."""

    CONTENT = """THE FUNDAMENTALS OF QUANTUM MECHANICS

Quantum mechanics is a fundamental theory in physics that provides a description of the 
physical properties of nature at the scale of atoms and subatomic particles. It is the 
foundation of all quantum physics including quantum chemistry, quantum field theory, 
quantum technology, and quantum information science.

Key Principles:

1. Wave-Particle Duality
Light and matter exhibit both wave-like and particle-like properties. The double-slit 
experiment demonstrates that particles like electrons can display interference patterns 
typical of waves when not observed, but behave as particles when measured.

2. Quantum Superposition
A quantum system can exist in multiple states simultaneously until it is measured. 
This is famously illustrated by Schrödinger's cat thought experiment, where a cat in a 
box can be both alive and dead until observed.

3. Quantum Entanglement
Two or more particles can become entangled such that the quantum state of each particle 
cannot be described independently of the state of the others. When a measurement is made 
on one entangled particle, it instantaneously affects the state of the other, regardless 
of the distance between them.

4. Heisenberg Uncertainty Principle
It is impossible to simultaneously know both the exact position and momentum of a 
particle. This is not a limitation of measurement technology, but a fundamental property 
of nature.

Applications:
- Quantum computers use superposition and entanglement to perform computations
- Quantum cryptography provides theoretically unbreakable encryption
- MRI machines rely on quantum mechanical principles
- Semiconductor electronics depend on quantum mechanics for their operation

The development of quantum mechanics in the early 20th century revolutionized physics 
and continues to drive technological innovation today.
"""


class HistoryTextFile:
    """Generate history-related text content."""

    CONTENT = """THE RENAISSANCE: A CULTURAL REBIRTH

The Renaissance was a cultural movement that profoundly affected European intellectual 
life in the early modern period. Beginning in Italy, it spread to the rest of Europe 
by the 16th century, influencing art, architecture, philosophy, and science.

Historical Context:

The Renaissance emerged from the Late Middle Ages, catalyzed by the Black Death pandemic 
that devastated Europe between 1347 and 1351. The resulting social, economic, and 
political changes created conditions favorable for cultural renewal.

Key Characteristics:

1. Humanism
Intellectual movement emphasizing the study of classical Greek and Roman texts. 
Humanists believed in the dignity and potential of individual humans, shifting focus 
from purely religious concerns to earthly achievements.

2. Artistic Innovation
Artists like Leonardo da Vinci, Michelangelo, and Raphael developed techniques 
including perspective, chiaroscuro, and anatomical accuracy. Art became more 
naturalistic and focused on human emotion and experience.

3. Scientific Advancement
The period saw breakthroughs in astronomy (Copernicus, Galileo), anatomy (Vesalius), 
and engineering (Da Vinci's designs). The scientific method began to take shape.

Major Figures:

- Leonardo da Vinci (1452-1519): Painter, sculptor, architect, scientist
- Michelangelo Buonarroti (1475-1564): Painter, sculptor, poet
- Niccolò Machiavelli (1469-1527): Political philosopher
- Gutenberg (c. 1398-1468): Inventor of the printing press

Impact:

The Renaissance laid the groundwork for the Protestant Reformation, the Scientific 
Revolution, and the Enlightenment. Its emphasis on human potential and empirical 
observation continues to influence modern thought and culture.
"""


class GeographyTextFile:
    """Generate geography-related text content."""

    CONTENT = """WORLD'S MAJOR BIOGENOMIC REGIONS AND ECOSYSTEMS

Earth's biodiversity is distributed across distinct biogeographic regions, each 
characterized by unique climate, flora, and fauna shaped by millions of years of 
evolution.

The Six Major Biogeographic Realms:

1. Palearctic Region
Covers Europe, Northern Asia, and North Africa. Features include:
- Temperate forests in Western Europe
- Tundra in Siberia
- Deserts in North Africa and Central Asia
- Iconic species: Brown bear, European bison, Siberian tiger

2. Nearctic Region
North America north of Mexico. Characterized by:
- Arctic tundra in Canada and Alaska
- Temperate rainforests in Pacific Northwest
- Grasslands (prairies) in central regions
- Iconic species: Bison, bald eagle, grizzly bear

3. Neotropical Region
Central and South America, including the Amazon rainforest:
- Largest rainforest ecosystem on Earth
- Highest species diversity per unit area
- Amazon River system
- Iconic species: Jaguar, toucan, sloth, poison dart frogs

4. Afrotropical Region
Sub-Saharan Africa:
- Savanna ecosystems (Serengeti)
- Congo Rainforest
- Desert environments (Sahara, Kalahari)
- Iconic species: African elephant, lion, giraffe, gorilla

5. Indomalayan Region
South and Southeast Asia:
- Tropical rainforests of Southeast Asia
- Monsoon forests of India
- Himalayan ecosystems
- Iconic species: Bengal tiger, orangutan, Asian elephant

6. Australasian Region
Australia, New Zealand, and Pacific Islands:
- Unique marsupial fauna
- Great Barrier Reef
- Eucalyptus forests
- Iconic species: Kangaroo, koala, platypus, kiwi bird

Conservation Challenges:

Climate change, habitat destruction, and poaching threaten these ecosystems. 
Understanding biogeographic patterns is crucial for conservation planning and 
biodiversity preservation efforts worldwide.
"""


class TechnologyTextFile:
    """Generate technology-related text content."""

    CONTENT = """ARTIFICIAL INTELLIGENCE: FROM THEORY TO APPLICATION

Artificial Intelligence (AI) has evolved from theoretical concept to transformative 
technology, reshaping industries and daily life. Understanding AI requires examining 
its foundations, evolution, and current applications.

Foundations of AI:

Artificial intelligence refers to computer systems capable of performing tasks that 
typically require human intelligence. These include:

- Learning and adaptation
- Reasoning and problem-solving
- Perception and interpretation
- Language understanding and generation
- Decision making

Key AI Paradigms:

1. Machine Learning (ML)
Systems that learn from data without explicit programming. Types include:
- Supervised Learning: Training with labeled data
- Unsupervised Learning: Finding patterns in unlabeled data
- Reinforcement Learning: Learning through reward-based feedback

2. Deep Learning
Neural networks with multiple layers that learn hierarchical representations:
- Convolutional Neural Networks (CNNs): Image recognition
- Recurrent Neural Networks (RNNs): Sequential data
- Transformers: Natural language processing

3. Large Language Models (LLMs)
Modern AI systems trained on vast text corpora:
- GPT (Generative Pre-trained Transformer) series
- BERT and its variants
- Specialized domain models

Current Applications:

- Healthcare: Drug discovery, medical imaging analysis, patient monitoring
- Finance: Fraud detection, algorithmic trading, risk assessment
- Transportation: Autonomous vehicles, route optimization
- Entertainment: Content recommendation, game AI
- Education: Personalized learning, automated grading

Ethical Considerations:

AI raises important questions about:
- Bias in AI systems and decision-making
- Privacy and surveillance
- Job displacement and economic impact
- Autonomous weapons and military applications
- AI alignment and safety

The future of AI depends on responsible development practices, transparency, and 
careful consideration of societal impacts. As AI continues to advance, collaborative 
efforts between technologists, policymakers, and ethicists will be essential to 
ensure beneficial outcomes for humanity."""


class SpaceTextFile:
    """Generate space-related text content."""

    CONTENT = """EXOPLANETS: SEARCHING FOR LIFE BEYOND EARTH

The discovery of exoplanets, which are planets orbiting stars outside our solar system, 
has revolutionized our understanding of the universe and the possibility of life beyond 
Earth. Astronomers have now confirmed over 5,500 exoplanets, with many more candidates 
awaiting confirmation.

History of Exoplanet Discovery:

1992: First exoplanets discovered orbiting pulsar PSR B1257+12
1995: First exoplanet discovered around a Sun-like star (51 Pegasi b)
2009: Kepler Space Telescope launched to find Earth-like planets
2015: Discovery of Proxima Centauri b, the closest potentially habitable exoplanet
2023: Over 5,500 confirmed exoplanets in our galaxy

Detection Methods:

1. Transit Method (Kepler/TESS Spacecraft)
The transit method measures periodic dips in starlight when a planet passes in front 
of its star from our viewpoint. The amount of light blocked reveals the planet size, 
while the orbital period reveals the distance from the star. NASA's Kepler mission 
found thousands of candidates using this method.

2. Radial Velocity Method
This technique detects the wobble of a star caused by the gravitational pull of 
orbiting planets. As a planet orbits, it causes the star to move slightly, creating 
a Doppler shift in the star's spectrum. This method is excellent for detecting massive 
planets close to their stars.

3. Direct Imaging
Taking actual pictures of exoplanets is extremely challenging due to the brightness 
difference between planets and their host stars. Astronomers use coronagraphs and 
extreme adaptive optics to block starlight and capture images of young, massive 
planets orbiting far from their stars.

The Habitable Zone:

The habitable zone, also known as the Goldilocks zone, is the region around a star 
where liquid water can exist on a planet's surface. This zone is not too hot and 
not too cold, but just right for water to remain in liquid form. Planets in this 
zone are prime candidates for hosting life as we know it.

Notable Exoplanets:

Proxima Centauri b: The closest known exoplanet to Earth, located just 4.2 light-years 
away in the habitable zone of Proxima Centauri.

TRAPPIST-1 System: Seven Earth-sized planets, several in the habitable zone, located 
about 40 light-years away.

Kepler-452b: Sometimes called Earth's cousin, this planet orbits a Sun-like star 
and is about 1.6 times the size of Earth.

Future Exploration:

The James Webb Space Telescope (JWST) is analyzing exoplanet atmospheres for signs 
of life-indicating molecules like oxygen, methane, and ozone. Future missions like 
the Habitable Worlds Observatory will directly image Earth-like planets and search 
for biosignatures that could indicate the presence of life.

The search for extraterrestrial life represents one of humanity's greatest scientific 
quests, with exoplanet research at the frontier of discovery."""


class PDFGenerator(FPDF):
    """Base class for PDF generation."""

    def header(self):
        self.set_font("helvetica", "B", 15)
        self.cell(0, 10, "RAG System Document", 0, 1, "C")
        self.ln(5)

    def chapter_title(self, title: str):
        self.set_font("helvetica", "B", 12)
        self.cell(0, 10, title, 0, 1, "L")
        self.ln(2)

    def chapter_body(self, body: str):
        self.set_font("helvetica", "", 10)
        self.multi_cell(0, 5, body)
        self.ln()


class RenewableEnergyPDF:
    """Generate PDF about renewable energy."""

    CONTENT = """
Renewable energy comes from natural sources that are constantly replenished. Unlike 
fossil fuels, renewable energy sources produce little to no greenhouse gas emissions 
and are key to combating climate change.

SOLAR ENERGY

Solar energy harnesses the power of the sun to generate electricity or heat. Solar 
panels convert sunlight into electricity using photovoltaic cells. Concentrated solar 
power uses mirrors to focus sunlight and generate heat for electricity generation. 
Solar energy is abundant, sustainable, and can be deployed at various scales from 
rooftop panels to large solar farms. Countries like China, United States, and India 
are leading in solar energy adoption.

WIND ENERGY

Wind turbines convert the kinetic energy of wind into electrical energy. Onshore 
wind farms are common in rural areas with consistent wind patterns. Offshore wind 
farms, located in bodies of water, can generate more power due to stronger and more 
consistent winds. Wind energy is one of the fastest-growing renewable energy sources 
worldwide, with Germany and China being major producers.

HYDROPOWER

Hydropower generates electricity by using the force of flowing or falling water to 
spin turbines. Dam-based hydropower provides reliable baseload power and can store 
water for later use during peak demand. Run-of-river hydropower uses the natural 
flow of rivers without large dams. Pumped storage facilities store energy by moving 
water between reservoirs at different elevations, acting as giant batteries.

GEOTHERMAL ENERGY

Geothermal energy utilizes heat from beneath the Earth's surface to generate 
electricity or provide direct heating. Geothermal power plants tap into underground 
reservoirs of steam or hot water. Countries like Iceland generate significant 
portions of their electricity from geothermal sources. Direct use applications 
include heating buildings, greenhouses, and industrial processes.

BIOMASS ENERGY

Biomass energy comes from organic materials including wood, agricultural residues, 
and dedicated energy crops. Biofuels like ethanol and biodiesel provide alternatives 
to fossil fuels for transportation. Biomass can be burned directly for heat or 
converted into biogas through anaerobic digestion. Sustainable biomass production 
is crucial to avoid competing with food production for land and resources.
"""


class ClimateSciencePDF:
    """Generate PDF about climate science."""

    CONTENT = """
Climate science studies long-term weather patterns and changes in Earth's climate 
system. Understanding climate dynamics is essential for predicting future changes 
and developing adaptation strategies to protect communities and ecosystems.

THE GREENHOUSE EFFECT

The greenhouse effect is a natural process that warms Earth's surface. Atmospheric 
gases like carbon dioxide, methane, and water vapor trap heat radiated from Earth's 
surface. Without the greenhouse effect, Earth would be about 33 degrees Celsius 
colder. However, human activities have intensified the greenhouse effect by burning 
fossil fuels and deforestation, leading to global warming.

CLIMATE FEEDBACK LOOPS

Climate systems contain feedback loops that can amplify or reduce changes:

Positive Feedbacks:
- Ice-albedo feedback: Melting ice exposes darker surfaces that absorb more heat
- Permafrost thawing releases stored methane, a potent greenhouse gas
- Water vapor feedback: Warmer air holds more water vapor, increasing warming

Negative Feedbacks:
- Increased plant growth absorbs more CO2 from the atmosphere
- Chemical weathering of rocks removes CO2 over geological timescales
- Ocean carbon uptake increases as temperatures rise

OCEAN CIRCULATION

Oceans absorb heat and carbon dioxide, moderating climate change. Ocean currents 
distribute heat around the planet, affecting weather patterns. The thermohaline 
circulation, also known as the Atlantic Meridional Overturning Circulation, affects 
climate in North America and Europe. Ocean acidification from absorbed CO2 threatens 
marine ecosystems and organisms that build calcium carbonate shells.

CLIMATE MODELS

Climate models simulate Earth's climate system using mathematical representations 
of atmosphere, ocean, land surface, and ice components. Models range from simple 
energy balance models to complex Earth system models. Model projections rely on 
scenarios of future greenhouse gas emissions. Models have successfully predicted 
many aspects of climate change, including global temperature rise and sea level increase.

IMPACTS AND ADAPTATION

Climate change affects ecosystems, weather extremes, and human societies. Sea level 
rise threatens coastal communities and infrastructure. More frequent extreme weather 
events cause damage and displacement. Changing precipitation patterns affect agriculture 
and water resources. Adaptation strategies include building sea walls, developing 
drought-resistant crops, and improving water management systems.
"""


class RoboticsPDF:
    """Generate PDF about robotics."""

    CONTENT = """
Robotics combines mechanical engineering, electrical engineering, and computer science 
to design, construct, and operate machines capable of performing tasks autonomously 
or with human guidance. Modern robotics has advanced significantly with improvements 
in AI, sensors, and actuators.

HISTORICAL DEVELOPMENT

The concept of mechanical humans dates back to ancient Greek mythology. The first 
programmable robot was the Unimate, invented in 1954 and used for manufacturing 
at General Motors. Industrial robots revolutionized assembly line production in the 
1960s and 1970s. Recent decades have seen exponential advances in artificial 
intelligence, sensors, and actuators.

ROBOT COMPONENTS

Actuators: Motors, hydraulics, or pneumatics that enable movement and provide force.

Sensors: Cameras, lidar, ultrasonic sensors, and encoders for perception of the 
environment and robot state.

Controllers: Processors that run control algorithms and make decisions based on 
sensor input and programmed instructions.

Effectors: Grippers, tools, and end-effectors for physically interacting with 
objects and the environment.

ROBOT CLASSIFICATION

Industrial Robots: Articulated robots with multiple rotating joints for flexibility. 
SCARA robots for horizontal movement and pick-and-place operations. Delta robots 
for high-speed pick and place. Cartesian robots for linear motion applications.

Service Robots: Medical robots for surgery and rehabilitation. Agricultural robots 
for planting, harvesting, and monitoring crops. Domestic robots for cleaning and 
assistance. Military robots for reconnaissance and logistics.

Autonomous Vehicles: Self-driving cars use sensors and AI for navigation. Drones 
provide aerial surveillance and delivery. Underwater robots explore ocean depths. 
Space robots maintain satellites and stations.

ARTIFICIAL INTELLIGENCE IN ROBOTICS

Modern robots incorporate AI for computer vision, natural language processing, 
reinforcement learning, simultaneous localization and mapping (SLAM), and predictive 
maintenance. These capabilities enable robots to operate in complex, unstructured 
environments and learn from experience.

ETHICS AND FUTURE

Robotics raises important ethical questions about automation and job displacement, 
safety standards for human-robot collaboration, legal liability for autonomous robot 
actions, privacy implications of surveillance robots, and the weaponization of robotic 
systems. The future promises continued innovation with robots becoming more capable, 
intelligent, and integrated into daily life.
"""


class DOCXGenerator:
    """Base class for DOCX generation."""

    @staticmethod
    def add_heading(doc: Document, text: str, level: int = 1):
        doc.add_heading(text, level)

    @staticmethod
    def add_paragraph(doc: Document, text: str):
        doc.add_paragraph(text)


class InternetHistoryDOCX:
    """Generate DOCX about internet history."""

    CONTENT = [
        ("The Evolution of the Internet", 1),
        ("", 0),
        (
            "The internet has transformed from a military research project to a global communication network that connects billions of people and powers the modern economy.",
            0,
        ),
        ("", 0),
        ("ARPANET: The Beginning", 1),
        (
            "The precursor to the internet was ARPANET, created by the U.S. Department of Defense Advanced Research Projects Agency in 1969. The first message was sent between UCLA and Stanford Research Institute on October 29, 1969.",
            0,
        ),
        ("", 0),
        ("Key Milestones", 1),
        ("1971: First email sent by Ray Tomlinson", 0),
        ("1983: TCP/IP becomes the standard protocol", 0),
        ("1989: Tim Berners-Lee invents the World Wide Web", 0),
        ("1993: Mosaic browser brings web to the masses", 0),
        ("1998: Google founded, revolutionizing search", 0),
        ("2004: Facebook launches social networking era", 0),
        ("", 0),
        ("Technical Architecture", 1),
        ("The internet operates through a layered architecture:", 0),
        ("Physical Layer: Fiber optic cables, wireless connections", 0),
        ("Data Link Layer: Ethernet, Wi-Fi protocols", 0),
        ("Network Layer: IP addressing and routing", 0),
        ("Transport Layer: TCP/UDP for data delivery", 0),
        ("Application Layer: HTTP, SMTP, FTP protocols", 0),
        ("", 0),
        ("Modern Internet Scale", 1),
        (
            "Today, the internet connects over 5 billion people worldwide. Data centers process exabytes of information daily. Cloud computing has democratized access to computing resources. The Internet of Things is connecting billions of devices from home appliances to industrial sensors.",
            0,
        ),
    ]


class DatabaseSystemsDOCX:
    """Generate DOCX about database systems."""

    CONTENT = [
        ("Database Management Systems", 1),
        ("", 0),
        (
            "Database management systems (DBMS) are software applications that enable users to define, create, maintain, and control access to databases. They serve as the foundation for modern data-driven applications.",
            0,
        ),
        ("", 0),
        ("Types of Database Models", 1),
        ("", 0),
        ("Relational Databases", 1),
        (
            "Relational databases organize data into tables with rows and columns. SQL (Structured Query Language) is used for querying and manipulation. Examples include PostgreSQL, MySQL, and Oracle. ACID properties ensure reliable transactions.",
            0,
        ),
        ("", 0),
        ("NoSQL Databases", 1),
        (
            "NoSQL databases provide flexible schemas and horizontal scalability. Document stores like MongoDB use JSON-like documents. Key-value stores like Redis offer fast lookups. Column-family stores like Cassandra handle massive datasets. Graph databases like Neo4j model relationships efficiently.",
            0,
        ),
        ("", 0),
        ("Key Concepts", 1),
        ("Normalization: Organizing data to reduce redundancy", 0),
        ("Indexing: Creating data structures for fast retrieval", 0),
        ("Transactions: Ensuring atomic, consistent operations", 0),
        ("Sharding: Distributing data across multiple servers", 0),
        ("Replication: Creating copies for fault tolerance", 0),
        ("", 0),
        ("Modern Trends", 1),
        (
            "Cloud databases have shifted infrastructure to managed services. Multi-model databases support multiple data models in one system. Time-series databases optimize for timestamped data. Vector databases enable similarity search for AI applications.",
            0,
        ),
    ]


class HTMLGenerator:
    """Base class for HTML generation."""

    @staticmethod
    def create_structure(title: str, heading: str, sections: list) -> str:
        sections_html = ""
        for section in sections:
            sections_html += f"""
        <section>
            <h2>{section["title"]}</h2>
            <p>{section["content"]}</p>
        </section>
"""
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        header {{
            background-color: #2c3e50;
            color: white;
            padding: 20px;
            text-align: center;
            border-radius: 5px;
        }}
        section {{
            background-color: white;
            padding: 20px;
            margin: 20px 0;
            border-radius: 5px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        h2 {{
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
    </style>
</head>
<body>
    <header>
        <h1>{heading}</h1>
    </header>
{sections_html}
</body>
</html>"""


class NeuralNetworksHTML:
    """Generate HTML about neural networks."""

    CONTENT = {
        "title": "Neural Networks: Fundamentals and Applications",
        "heading": "Understanding Neural Networks",
        "sections": [
            {
                "title": "What are Neural Networks?",
                "content": "Neural networks are computing systems inspired by biological neural networks in the brain. They consist of interconnected nodes (neurons) organized in layers. Input data flows through hidden layers, undergoing transformations, to produce outputs. Through training, networks learn to recognize patterns and make predictions.",
            },
            {
                "title": "Architecture Components",
                "content": "A typical neural network contains an input layer that receives data, hidden layers that perform computations, and an output layer that produces results. Each connection between neurons has a weight that is adjusted during training. Activation functions like ReLU, sigmoid, and tanh introduce non-linearity into the network.",
            },
            {
                "title": "Training Process",
                "content": "Neural networks learn through backpropagation, where errors are calculated and weights are adjusted. The gradient descent optimization algorithm minimizes the loss function. Training requires large datasets and significant computational resources. Techniques like dropout and regularization prevent overfitting.",
            },
            {
                "title": "Types of Neural Networks",
                "content": "Convolutional Neural Networks (CNNs) excel at image recognition. Recurrent Neural Networks (RNNs) handle sequential data like text and time series. Transformer networks power modern language models. Generative Adversarial Networks (GANs) create realistic synthetic data.",
            },
            {
                "title": "Applications",
                "content": "Neural networks enable facial recognition, machine translation, voice assistants, autonomous vehicles, medical diagnosis, drug discovery, and many more applications. They have revolutionized computer vision and natural language processing.",
            },
        ],
    }


class BlockchainHTML:
    """Generate HTML about blockchain."""

    CONTENT = {
        "title": "Blockchain Technology: Beyond Cryptocurrency",
        "heading": "Introduction to Blockchain",
        "sections": [
            {
                "title": "What is Blockchain?",
                "content": "A blockchain is a distributed, immutable ledger that records transactions across multiple computers. Each block contains a cryptographic hash of the previous block, creating a chain. This structure makes it virtually impossible to alter past records without detection.",
            },
            {
                "title": "Key Concepts",
                "content": "Decentralization removes single points of failure by distributing control across network participants. Consensus mechanisms like Proof of Work and Proof of Stake validate transactions. Smart contracts are self-executing programs stored on the blockchain. Public keys and private keys enable secure ownership and transfers.",
            },
            {
                "title": "How Transactions Work",
                "content": "Users create transactions using their private keys. Transactions are broadcast to the network of nodes. Miners or validators group transactions into blocks. Consensus protocols verify the block validity. Once confirmed, the block is added to the blockchain permanently.",
            },
            {
                "title": "Beyond Cryptocurrency",
                "content": "Blockchain enables supply chain tracking, ensuring authenticity of products. Decentralized finance (DeFi) offers financial services without intermediaries. Non-fungible tokens (NFTs) represent ownership of digital assets. Blockchain can secure voting systems, land registries, and identity management.",
            },
            {
                "title": "Challenges and Future",
                "content": "Scalability remains a challenge as blockchain networks handle limited transactions per second. Energy consumption raises environmental concerns, especially in Proof of Work systems. Regulatory uncertainty creates business risks. Layer 2 solutions and new consensus mechanisms aim to address these limitations.",
            },
        ],
    }


def generate_txt_files(output_dir: Path) -> list[str]:
    """Generate 5 TXT files on general knowledge topics."""
    files = []

    txt_contents = [
        ("01_science.txt", ScienceTextFile.CONTENT),
        ("02_history.txt", HistoryTextFile.CONTENT),
        ("03_geography.txt", GeographyTextFile.CONTENT),
        ("04_technology.txt", TechnologyTextFile.CONTENT),
        ("05_space.txt", SpaceTextFile.CONTENT),
    ]

    for filename, content in txt_contents:
        filepath = output_dir / filename
        filepath.write_text(content, encoding="utf-8")
        files.append(str(filepath))
        print(f"  Created {filename}")

    return files


def generate_pdf_files(output_dir: Path) -> list[str]:
    """Generate 3 PDF files using fpdf2."""
    files = []

    pdf_contents = [
        ("06_renewable_energy.pdf", RenewableEnergyPDF.CONTENT),
        ("07_climate_science.pdf", ClimateSciencePDF.CONTENT),
        ("08_robotics.pdf", RoboticsPDF.CONTENT),
    ]

    for filename, content in pdf_contents:
        pdf = PDFGenerator()
        pdf.add_page()
        pdf.chapter_title(filename.replace(".pdf", "").replace("_", " ").title())
        pdf.chapter_body(content)

        filepath = output_dir / filename
        pdf.output(str(filepath))
        files.append(str(filepath))
        print(f"  Created {filename}")

    return files


def generate_docx_files(output_dir: Path) -> list[str]:
    """Generate 2 DOCX files using python-docx."""
    files = []

    docx_contents = [
        ("09_internet_history.docx", InternetHistoryDOCX.CONTENT),
        ("10_database_systems.docx", DatabaseSystemsDOCX.CONTENT),
    ]

    for filename, content in docx_contents:
        doc = Document()
        for item in content:
            if item[1] == 1:
                doc.add_heading(item[0], level=1)
            elif item[1] == 0:
                if item[0]:
                    doc.add_paragraph(item[0])

        filepath = output_dir / filename
        doc.save(str(filepath))
        files.append(str(filepath))
        print(f"  Created {filename}")

    return files


def generate_html_files(output_dir: Path) -> list[str]:
    """Generate 2 HTML files with structured content."""
    files = []

    html_contents = [
        ("11_neural_networks.html", NeuralNetworksHTML.CONTENT),
        ("12_blockchain.html", BlockchainHTML.CONTENT),
    ]

    for filename, content in html_contents:
        html = HTMLGenerator.create_structure(
            content["title"], content["heading"], content["sections"]
        )

        filepath = output_dir / filename
        filepath.write_text(html, encoding="utf-8")
        files.append(str(filepath))
        print(f"  Created {filename}")

    return files


def main():
    """Generate all dummy data files."""
    output_dir = Path(__file__).parent / "raw"
    output_dir.mkdir(parents=True, exist_ok=True)

    print("Generating TXT files...")
    txt_files = generate_txt_files(output_dir)
    print(f"Created {len(txt_files)} TXT files")

    print("\nGenerating PDF files...")
    pdf_files = generate_pdf_files(output_dir)
    print(f"Created {len(pdf_files)} PDF files")

    print("\nGenerating DOCX files...")
    docx_files = generate_docx_files(output_dir)
    print(f"Created {len(docx_files)} DOCX files")

    print("\nGenerating HTML files...")
    html_files = generate_html_files(output_dir)
    print(f"Created {len(html_files)} HTML files")

    all_files = txt_files + pdf_files + docx_files + html_files
    print(f"\nTotal: {len(all_files)} files generated in {output_dir}")


if __name__ == "__main__":
    main()
